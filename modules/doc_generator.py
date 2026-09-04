from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

def create_pdd_docx(sections):
    """
    Generates a styled Word document (.docx) from a dictionary of sections
    and returns a BytesIO buffer suitable for Streamlit downloading.
    """
    doc = Document()
    
    # ---------------------------------
    # STYLING CONFIG (Subtle & Professional)
    # ---------------------------------
    # Style default normal text
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    
    # Add Document Title
    title = doc.add_paragraph()
    title_run = title.add_run("Process Definition Document (PDD)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Slate Blue
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Generated automatically by AI Automation Documentation Generator")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F) # Gray
    
    doc.add_paragraph("\n") # spacing
    
    # Loop over the sections
    for section_title, content in sections.items():
        # Add heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title)
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Slate Blue
        
        # Add a subtle bottom border or spacing
        # For simplicity, we just add paragraphs
        lines = content.split('\n')
        for line in lines:
            trimmed = line.strip()
            if not trimmed:
                continue
            
            # Handle headers inside markdown (e.g. ### Subheading)
            if trimmed.startswith('###'):
                p = doc.add_paragraph()
                r = p.add_run(trimmed.replace('###', '').strip())
                r.font.bold = True
                r.font.size = Pt(12)
            elif trimmed.startswith('##'):
                p = doc.add_paragraph()
                r = p.add_run(trimmed.replace('##', '').strip())
                r.font.bold = True
                r.font.size = Pt(13)
            elif trimmed.startswith('#'):
                p = doc.add_paragraph()
                r = p.add_run(trimmed.replace('#', '').strip())
                r.font.bold = True
                r.font.size = Pt(14)
            # Handle bullet lists
            elif trimmed.startswith('* ') or trimmed.startswith('- '):
                bullet_content = trimmed[2:].strip()
                doc.add_paragraph(bullet_content, style='List Bullet')
            # Handle numbered lists
            elif trimmed.startswith('1. ') or trimmed.startswith('2. ') or trimmed.startswith('3. ') or trimmed.startswith('4. ') or trimmed.startswith('5. ') or trimmed.startswith('6. ') or trimmed.startswith('7. ') or trimmed.startswith('8. ') or trimmed.startswith('9. '):
                # find first dot index
                dot_idx = trimmed.find('.')
                list_content = trimmed[dot_idx+1:].strip()
                doc.add_paragraph(list_content, style='List Number')
            else:
                doc.add_paragraph(trimmed)
                
        doc.add_paragraph("\n") # space between sections
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
